import os
import json
import tempfile
from typing import List, Optional, Dict, Union

import streamlit as st
from streamlit_option_menu import option_menu
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx.shared import Inches

# Configuration Constants
SUBMISSIONS_FILE = "submissions.json"
MAX_FILE_SIZE_MB = 10
ALLOWED_IMAGE_TYPES = ["png", "jpg", "jpeg"]
BASE_PATH = "/home/fortinet/streamlit_env/TPB"

# Technologies and Models Lists
FORTINET_TECHNOLOGIES = [
    'NGFW', 'FMG', 'FAZ', 'FSA', 'SOAR-MT-HA', 'SOAR-ENT-Single', 
    'SIEM', 'SDWAN', 'FNDR', 'FNAC', 'FDC', 'FWB', 'FAD', 'FSW', 
    'FAP', 'FPAM', 'FEX', 'FML', 'SASE', 'FAC', 'FortiRecon', 
    'FortiTester', 'EDR', 'FortiMonitor', 'FPX'
]

NGFW_MODELS = [
    'FG-4201F', 'FG-3001F', 'FG-2601F', 'FG-1101E', 'FG-1001F', 
    'FG-901G', 'FG-601F', 'FG-401F', 'FG-201F', 'FG-121G', 
    'FG-101F', 'FG-91G', 'FG-40F', 'FGR-70F-3G-4G', 'Azure-FG-VM16', 
    'Azure-FG-VM04', 'FNAC-500F', 'FAZ-800G', 'FAZ-150G', 'FAZ-VM-GB200', 
    'FMG-VM', 'FMG-200G', 'FTS-2000F', 'FEX-511F', 'FSW-1048E', 
    'FSW-448E', 'FSW-424E', 'FSW-108F', 'FSA-VM00', 'FSA-500G', 
    'FAC-VM', 'FPAM-VM', 'FAI-VM16', 'FWB-VM04', 'SOAR-VM', 
    'EDR-on', 'FPX-400G', 'FML-400F', 'FAP-231G', 'FAP-233G', 
    'FAP-234G', 'FAD-220F', 'FAD-320F', 'FAD-420F', 
    'FAD-1200F', 'FAD-2200F', 'FAD-4200F'
]

def custom_css():
    """Apply custom CSS styling."""
    st.markdown("""
    <style>
    body {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .container {
        max-width: 800px;
        margin: auto;
        padding: 20px;
        background-color: #f9f9f9;
        border-radius: 10px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
    }
    h1 {
        color: #0078d7;
        text-align: center;
        font-size: 2.7em;
        margin-bottom: 20px;
    }
    .stButton > button {
        background-color: #007bff;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .stButton > button:hover {
        background-color: #0056b3;
    }
    </style>
    """, unsafe_allow_html=True)

def validate_input(name: str, proj: str) -> bool:
    """Validate input fields."""
    if not name or not proj:
        st.error("Customer Name and Project Name are required!")
        return False
    return True

def save_temp_file(uploaded_file):
    """Save uploaded file to a temporary location."""
    if uploaded_file is None:
        return None
    
    if uploaded_file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        st.error(f"File size exceeds {MAX_FILE_SIZE_MB} MB")
        return None
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmpfile:
        tmpfile.write(uploaded_file.getbuffer())
        return tmpfile.name

def replace_placeholder(doc, placeholder, replacement):
    """Replace placeholders in document."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, replacement)

def replace_image_placeholder(doc, placeholder, image_path, width):
    """Replace image placeholders in document."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_image_placeholder(cell, placeholder, image_path, width)

def save_submission(submission):
    """Save submission to JSON file."""
    try:
        submissions = load_submissions()
        submissions.append(submission)
        with open(SUBMISSIONS_FILE, 'w') as f:
            json.dump(submissions, f, indent=4)
    except Exception as e:
        st.error(f"Error saving submission: {e}")

def load_submissions():
    """Load submissions from JSON file."""
    try:
        if os.path.exists(SUBMISSIONS_FILE):
            with open(SUBMISSIONS_FILE, 'r') as f:
                return json.load(f)
        return []
    except Exception as e:
        st.error(f"Error loading submissions: {e}")
        return []

def create_proposal(name, proj, logo, topology, technologies, models, comments):
    """Create technical proposal document."""
    if not validate_input(name, proj):
        return None

    logo_path = save_temp_file(logo)
    topology_path = save_temp_file(topology)

    try:
        master = Document_compose(os.path.join(BASE_PATH, 'section0.docx'))
        
        if logo_path:
            replace_image_placeholder(master, "(Logo)", logo_path, 1.25)
            os.unlink(logo_path)

        replace_placeholder(master, "(Proj)", proj)
        replace_placeholder(master, "<<Customer Name>>", name)
        
        composer = Composer(master)
        
        models_paths = [os.path.join(BASE_PATH, 'Models', f"{model}.docx") for model in models]
        technologies_paths = [os.path.join(BASE_PATH, f"{tech}.docx") for tech in technologies]
        mid = [os.path.join(BASE_PATH, 'Design.docx')]
        
        combined_list = technologies_paths + mid + models_paths
        
        process_documents(composer, combined_list, name, topology_path, comments)
        
        filename = f"{name}_{proj}_Technical Proposal ver 1.0.docx"
        composer.save(os.path.join(BASE_PATH, filename))
        
        return filename
    
    except Exception as e:
        st.error(f"Error creating proposal: {e}")
        return None

def process_documents(composer, document_list, name, topology_path, comments):
    """Process and combine documents."""
    missing_files = []
    for item in document_list:
        try:
            doc2 = Document_compose(item)
            replace_placeholder(doc2, "<<Customer Name>>", name)
            
            if item.endswith("Design.docx"):
                if topology_path:
                    replace_image_placeholder(doc2, "<<Design>>", topology_path, 6)
                    os.unlink(topology_path)
                replace_placeholder(doc2, "<<Design Describtion>>", comments)
            
            composer.append(doc2)
        except Exception as e:
            missing_files.append(item)
            st.error(f"Error appending {item}: {e}")

def main():
    """Main Streamlit application."""
    custom_css()
    st.image("https://upload.wikimedia.org/wikipedia/commons/6/62/Fortinet_logo.svg", use_column_width=True)
    st.title("Fortinet Technical Proposal Builder")

    with st.sidebar:
        selected = option_menu("Main Menu", ["Create Proposal", 'Submitted Proposals'], 
            icons=['house', 'gear'], menu_icon="cast", default_index=0)

    if selected == "Create Proposal":
        with st.form(key='Fortinet Technical Proposal Survey'):
            name = st.text_input("Customer Name")
            proj = st.text_input("Project Name")
            logo = st.file_uploader("Upload Company Logo", type=ALLOWED_IMAGE_TYPES)
            topology = st.file_uploader("Upload Design Topology", type=ALLOWED_IMAGE_TYPES)
            
            technologies = st.multiselect(
                "Which Fortinet technologies are you offering?",
                FORTINET_TECHNOLOGIES
            )
            
            models = st.multiselect(
                "Which NGFW Hardware Models are you offering?",
                NGFW_MODELS
            )
            
            comments = st.text_area("Design Description")
            submit_button = st.form_submit_button(label='Submit')

        if submit_button:
            filename = create_proposal(name, proj, logo, topology, technologies, models, comments)
            
            if filename:
                st.success("Proposal created successfully!")
                with open(os.path.join(BASE_PATH, filename), "rb") as file:
                    st.download_button(
                        label="Download Technical Proposal",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                submission = {
                    "Customer Name": name,
                    "Project Name": proj,
                    "Technologies": technologies,
                    "Models": models,
                    "Comments": comments,
                    "Filename": filename
                }
                save_submission(submission)

    elif selected == "Submitted Proposals":
        st.subheader("Submitted Proposals")
        submissions = load_submissions()
        
        if submissions:
            for submission in submissions:
                st.markdown(f"""
                **Customer Name:** {submission['Customer Name']}  
                **Project Name:** {submission['Project Name']}  
                **Technologies:** {', '.join(submission['Technologies'])}  
                **Models:** {', '.join(submission['Models'])}  
                **Comments:** {submission['Comments']}
                """)
                
                with open(os.path.join(BASE_PATH, submission["Filename"]), "rb") as file:
                    st.download_button(
                        label="Download Proposal",
                        data=file,
                        file_name=submission["Filename"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.write("---")
        else:
            st.write("No submissions yet.")

    st.write("---")
    st.markdown("For feedback: [Hany Badawy](mailto:habouhaswa@fortinet.com)")
    st.write("© 2024 Fortinet UAE SE TEAM")

if __name__ == "__main__":
    main()
